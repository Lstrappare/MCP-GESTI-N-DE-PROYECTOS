from __future__ import annotations

from datetime import datetime
from pathlib import Path
from io import BytesIO
from typing import List, Optional
import re

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

from models import DocumentoEDTInput, TareaBase
from config.settings import (
    _COLOR_PRIMARIO,
    _COLOR_SECUNDARIO,
    _COLOR_ACENTO,
    _COLOR_TEXTO,
    _COLOR_GRIS_CLARO,
    _COLOR_BLANCO,
    _COLORES_ETAPA,
    _COLOR_SUBTAREA_BASE,
    _PREFIJO_ARCHIVO,
    _CARPETA_DOWNLOADS,
)
from services.mermaid_service import generar_imagen_mermaid_para_docx


# ==============================================================================
# HELPERS GENÉRICOS DRY — ESTILOS
# ==============================================================================

def aplicar_estilo_runs(paragraph, font_size=10, bold=False, color=None, font_name="Calibri"):
    """Aplica estilo a todos los runs de un párrafo."""
    for run in paragraph.runs:
        run.font.size = Pt(font_size)
        run.font.bold = bold
        run.font.name = font_name
        if color:
            run.font.color.rgb = color


def aplicar_estilo_celda(cell, fondo=None, borde=None, fuente_color=None, font_size=10, bold=False, alignment=None):
    """Aplica estilos completos a una celda de tabla."""
    if fondo:
        _set_cell_shading(cell, fondo)
    if borde:
        _set_cell_borders(
            cell, top=borde, bottom=borde, start=borde, end=borde,
        )
    for paragraph in cell.paragraphs:
        if alignment is not None:
            paragraph.alignment = alignment
        for run in paragraph.runs:
            run.font.size = Pt(font_size)
            run.font.bold = bold
            run.font.name = "Calibri"
            if fuente_color:
                run.font.color.rgb = fuente_color


def aplicar_bordes_uniformes(cell, size="4", color="BFBFBF", val="single"):
    """Aplica bordes uniformes a una celda."""
    border_style = {"sz": size, "val": val, "color": color}
    _set_cell_borders(
        cell, top=border_style, bottom=border_style,
        start=border_style, end=border_style,
    )


def crear_tabla_con_encabezados(doc, headers, colores, widths=None):
    """Crea una tabla con fila de encabezados estilizada."""
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    header_row = table.rows[0]
    for i, header in enumerate(headers):
        cell = header_row.cells[i]
        cell.text = header
        _set_cell_shading(cell, colores[i] if isinstance(colores, list) else colores)
        _aplicar_bordes_thick(cell)
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(10)
                run.font.name = "Calibri"
                run.font.color.rgb = _COLOR_BLANCO

    if widths:
        for i, width in enumerate(widths):
            header_row.cells[i].width = width

    return table


def _aplicar_bordes_thick(cell):
    border_style = {"sz": "6", "val": "single", "color": "0D2B4E"}
    _set_cell_borders(
        cell, top=border_style, bottom=border_style,
        start=border_style, end=border_style,
    )


def agregar_fila_estilizada(table, valores, fondo=None, borde=None, bold=False, font_size=10, fuente_color=None):
    """Agrega una fila a la tabla con estilos uniformes en todas sus celdas."""
    row = table.add_row()
    for i, valor in enumerate(valores):
        cell = row.cells[i]
        cell.text = str(valor)
        aplicar_estilo_celda(
            cell, fondo=fondo, borde=borde,
            fuente_color=fuente_color, font_size=font_size, bold=bold,
        )
    return row


# ==============================================================================
# HELPERS DE BAJO NIVEL — SHADING y BORDES
# ==============================================================================

def _set_cell_width(cell, width):
    """Establece el ancho de una celda."""
    cell.width = width


def _merge_horizontal(table, row, start_col, end_col):
    """Fusiona celdas horizontalmente en una fila."""
    for col in range(end_col, start_col, -1):
        table.cell(row, start_col).merge(table.cell(row, col))


def _merge_vertical(table, col, start_row, end_row):
    """Fusiona celdas verticalmente en una columna."""
    for row in range(end_row, start_row, -1):
        table.cell(start_row, col).merge(table.cell(row, col))


def _set_cell_text(cell, text, alignment=None, font_size=10, bold=False, color=None):
    """Establece el texto de una celda con formato."""
    for p in cell.paragraphs:
        p.clear()
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.name = "Calibri"
    if color:
        run.font.color.rgb = color
    if alignment is not None:
        p.alignment = alignment


def _aplicar_estilo_tabla_plantilla(table):
    """Aplica bordes sutiles y estilos por defecto a una tabla de plantilla."""
    border_style = {"sz": "4", "val": "single", "color": "BFBFBF"}
    for row in table.rows:
        for cell in row.cells:
            _set_cell_borders(
                cell, top=border_style, bottom=border_style,
                start=border_style, end=border_style,
            )
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    if run.font.size is None:
                        run.font.size = Pt(9)
                    if run.font.name is None:
                        run.font.name = "Calibri"
                    if run.font.color.rgb is None:
                        run.font.color.rgb = _COLOR_TEXTO

def _set_cell_shading(cell, color_hex: str) -> None:
    """Aplica color de fondo a una celda de tabla Word."""
    shading_elm = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>'
    )
    cell._tc.get_or_add_tcPr().append(shading_elm)


def _set_cell_borders(cell, **kwargs) -> None:
    """
    Aplica bordes a una celda. Parámetros: top, bottom, start, end.
    Cada uno recibe un dict con: sz (tamaño), val (tipo), color.
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')

    for edge, attrs in kwargs.items():
        element = parse_xml(
            f'<w:{edge} {nsdecls("w")} '
            f'w:sz="{attrs.get("sz", 4)}" '
            f'w:val="{attrs.get("val", "single")}" '
            f'w:color="{attrs.get("color", "BFBFBF")}"/>'
        )
        tcBorders.append(element)

    tcPr.append(tcBorders)


def _set_table_borders(table, size="4", color="000000"):
    """Aplica bordes a nivel de tabla (cubre celdas fusionadas correctamente)."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}/>')
        tbl.insert(0, tblPr)

    # Remover bordes existentes
    for existing in tblPr.findall(qn('w:tblBorders')):
        tblPr.remove(existing)

    borders_xml = (
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'  <w:left w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'  <w:bottom w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'  <w:right w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'  <w:insideH w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'  <w:insideV w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(parse_xml(borders_xml))


# ==============================================================================
# HELPERS — PÁRRAFOS
# ==============================================================================

def _add_styled_paragraph(doc, text, font_size=10, bold=False, color=None, alignment=None, space_after=6):
    """Agrega un párrafo con estilo al documento."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.name = "Calibri"
    if color:
        run.font.color.rgb = color
    if alignment is not None:
        p.alignment = alignment
    p.paragraph_format.space_after = Pt(space_after)
    return p


# ==============================================================================
# RECURSIVIDAD — TABLA BASE EDT
# ==============================================================================

def _limpiar_prefijo(nombre: str) -> str:
    """Remueve el prefijo numérico (ej. '1.1.1.1 ') del nombre de una tarea."""
    m = re.match(r"^[\d.]+\.?\s+", nombre)
    if m:
        return nombre[m.end():]
    return nombre


def _agregar_filas_tabla_recursivo(
    table, tarea: TareaBase, codigo_padre: str, indice: int,
    nivel: int, color_etapa_idx: int
) -> int:
    """
    Agrega filas recursivamente a la tabla de la EDT.
    Retorna el número total de tareas procesadas.
    """
    codigo = f"{codigo_padre}.{indice}"
    row = table.add_row()
    cells = row.cells

    nombre_limpio = _limpiar_prefijo(tarea.nombre)
    indent = "    " * (nivel - 2) if nivel >= 2 else ""
    nombre_display = f"{indent}{nombre_limpio}" if nivel >= 3 else nombre_limpio

    cells[0].text = codigo
    cells[1].text = nombre_display
    cells[2].text = tarea.descripcion_operativa or "N/A"
    cells[3].text = tarea.hito or "N/A"

    for cell in cells:
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in paragraph.runs:
                run.font.size = Pt(9)
                run.font.name = "Calibri"
                run.font.color.rgb = _COLOR_TEXTO

    if nivel == 1:
        color_hex = _COLORES_ETAPA[color_etapa_idx % len(_COLORES_ETAPA)]
        for cell in cells:
            _set_cell_shading(cell, color_hex)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(10)
    elif nivel == 2:
        for cell in cells:
            _set_cell_shading(cell, _COLOR_SUBTAREA_BASE)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(9)

    border_style = {"sz": "4", "val": "single", "color": "BFBFBF"}
    for cell in cells:
        _set_cell_borders(
            cell, top=border_style, bottom=border_style,
            start=border_style, end=border_style,
        )

    count = 1
    if tarea.subtareas:
        for sub_idx, subtarea in enumerate(tarea.subtareas, start=1):
            count += _agregar_filas_tabla_recursivo(
                table, subtarea, codigo, sub_idx,
                nivel + 1, color_etapa_idx,
            )

    return count


# ==============================================================================
# SECCIÓN 1 — ENCABEZADO PLANTILLA CORPORATIVA
# ==============================================================================

def _construir_tabla_firmas(doc) -> None:
    """Inserta una tabla de firmas con los 3 roles de aprobación."""
    _add_styled_paragraph(
        doc, "Firmas de Validación",
        font_size=12, bold=True, color=_COLOR_PRIMARIO,
        alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=8,
    )

    headers_firma = ["Rol", "Nombre Completo", "Firma"]
    tabla_firmas = crear_tabla_con_encabezados(
        doc, headers_firma, "0D2B4E",
    )

    for _ in range(3):
        tabla_firmas.add_row()

    roles = ["Director del proyecto", "Patrocinador del proyecto", "Cliente del proyecto"]
    for i, rol in enumerate(roles, start=1):
        tabla_firmas.cell(i, 0).text = rol
        tabla_firmas.cell(i, 1).text = ""
        tabla_firmas.cell(i, 2).text = ""

        for j in range(3):
            cell = tabla_firmas.cell(i, j)
            border_style = {"sz": "4", "val": "single", "color": "A9C4D8"}
            _set_cell_borders(
                cell, top=border_style, bottom=border_style,
                start=border_style, end=border_style,
            )
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)
                    run.font.name = "Calibri"

            if j == 2:
                cell.width = Cm(5)

        for paragraph in tabla_firmas.cell(i, 0).paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = _COLOR_PRIMARIO

    for row in tabla_firmas.rows:
        row.cells[0].width = Cm(5)
        row.cells[1].width = Cm(7)
        row.cells[2].width = Cm(5)

    doc.add_paragraph()

def _construir_encabezado_plantilla(doc, datos: DocumentoEDTInput, fecha: str, tipo_encabezado: str = "EDT") -> None:
    """Construye el encabezado corporativo estilo tablero."""
    if tipo_encabezado == "Minuta":
        texto_tipo_doc = "Nombre Minuta:\nEstructura de Desglose de Trabajo (EDT)"
        texto_numero = "Minuta No.:\n8"
    else:
        texto_tipo_doc = "Nombre Plantilla:\nEstructura de Desglose de Trabajo (EDT)"
        texto_numero = "Plantilla No.:\n8"

    # ── Logo + Empresa fuera de la tabla ──
    p_logo = doc.add_paragraph()
    p_logo.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_icon = p_logo.add_run("◼ ")
    run_icon.font.size = Pt(14)
    run_icon.font.bold = True
    run_icon.font.name = "Calibri"
    run_icon.font.color.rgb = _COLOR_PRIMARIO
    run_nombre = p_logo.add_run(datos.nombre_empresa or "MINTRANET")
    run_nombre.font.size = Pt(14)
    run_nombre.font.bold = True
    run_nombre.font.name = "Calibri"
    run_nombre.font.color.rgb = _COLOR_PRIMARIO

    # ── Tabla 4 columnas × 3 filas ──
    table = doc.add_table(rows=3, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    # Columna 0: vertical merge rows 0-2 → Logo + Empresa
    _merge_vertical(table, 0, 0, 2)
    cell_logo = table.cell(0, 0)
    _set_cell_text(
        cell_logo, f"◼ LOGO\n\n{datos.nombre_empresa or 'MINTRANET'}",
        alignment=WD_ALIGN_PARAGRAPH.CENTER, font_size=10, bold=True,
        color=_COLOR_PRIMARIO,
    )
    _set_cell_width(cell_logo, Cm(3.5))

    # Fila 0, columnas 1-2: horizontal merge → Nombre tipo de documento
    _merge_horizontal(table, 0, 1, 2)
    cell_tipo_doc = table.cell(0, 1)
    _set_cell_text(
        cell_tipo_doc, texto_tipo_doc,
        alignment=WD_ALIGN_PARAGRAPH.CENTER, font_size=10, bold=False,
        color=_COLOR_TEXTO,
    )

    # Fila 0, columna 3: Etapa
    cell_etapa = table.cell(0, 3)
    _set_cell_text(
        cell_etapa, "Etapa:\n1) Planeación",
        alignment=WD_ALIGN_PARAGRAPH.CENTER, font_size=9, bold=False,
        color=_COLOR_TEXTO,
    )
    _set_cell_width(cell_etapa, Cm(3))

    # Columna 1, filas 1-2: merge vertical → Título del proyecto + ID
    _merge_vertical(table, 1, 1, 2)
    cell_titulo = table.cell(1, 1)
    _set_cell_text(
        cell_titulo,
        f"Título del proyecto:\n{datos.nombre_proyecto}\n\nID del Proyecto:\n{datos.id_proyecto or 'N/A'}",
        alignment=WD_ALIGN_PARAGRAPH.CENTER, font_size=9, bold=False,
        color=_COLOR_TEXTO,
    )
    _set_cell_width(cell_titulo, Cm(5.5))

    # Columna 2, filas 1-2: merge vertical → Presupuestos
    _merge_vertical(table, 2, 1, 2)
    cell_presup = table.cell(1, 2)
    _set_cell_text(
        cell_presup,
        f"Presupuesto:\n{datos.presupuesto_fase or 'N/A'}\n\nPresupuesto del Proyecto:\n{datos.presupuesto_proyecto or 'N/A'}",
        alignment=WD_ALIGN_PARAGRAPH.CENTER, font_size=9, bold=False,
        color=_COLOR_TEXTO,
    )
    _set_cell_width(cell_presup, Cm(5.5))

    # Columna 3, fila 1: Número de plantilla/minuta
    cell_numero = table.cell(1, 3)
    _set_cell_text(
        cell_numero, texto_numero,
        alignment=WD_ALIGN_PARAGRAPH.CENTER, font_size=9, bold=False,
        color=_COLOR_TEXTO,
    )

    # Columna 3, fila 2: Fecha
    cell_fecha = table.cell(2, 3)
    _set_cell_text(
        cell_fecha, f"Fecha:\n{fecha}",
        alignment=WD_ALIGN_PARAGRAPH.CENTER, font_size=9, bold=False,
        color=_COLOR_TEXTO,
    )

    # Aplicar bordes negros a nivel tabla (cubre celdas fusionadas)
    _set_table_borders(table, size="4", color="000000")

    doc.add_paragraph()


# ==============================================================================
# SECCIÓN 2 — DIAGRAMA EDT
# ==============================================================================

def _construir_seccion_diagrama(doc, image_bytes: bytes) -> None:
    """Inserta el diagrama EDT como imagen centrada."""
    image_stream = BytesIO(image_bytes)
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_img = p_img.add_run()
    run_img.add_picture(image_stream, width=Inches(6.5))
    doc.add_paragraph()


# ==============================================================================
# SECCIÓN 3 — TABLA BASE DE LA EDT
# ==============================================================================

def _construir_seccion_tabla(doc, datos: DocumentoEDTInput) -> None:
    """Construye la Tabla Base de la EDT jerarquizada."""
    _add_styled_paragraph(
        doc, "TABLA BASE DE LA EDT",
        font_size=16, bold=True, color=_COLOR_PRIMARIO,
        alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=6,
    )

    p_sep = doc.add_paragraph()
    run_sep = p_sep.add_run("─" * 80)
    run_sep.font.color.rgb = _COLOR_GRIS_CLARO
    run_sep.font.size = Pt(8)

    _add_styled_paragraph(
        doc,
        "La siguiente tabla presenta la codificación EDT jerárquica, la descripción "
        "operativa de cada tarea y sus hitos asociados.",
        font_size=10, bold=False, color=_COLOR_TEXTO,
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=12,
    )

    headers = [
        "Nivel EDT",
        "Fase / Área / Actividad",
        "Relación con el Plan Operativo y Plan General de Trabajo",
        "Hito del Acta de Constitución (Cronograma)",
    ]
    table = crear_tabla_con_encabezados(
        doc, headers, "0D2B4E",
        widths=[Cm(2.5), Cm(4.5), Cm(6.5), Cm(4)],
    )

    # Fila del proyecto raíz
    root_row = table.add_row()
    root_row.cells[0].text = "0"
    root_row.cells[1].text = datos.nombre_proyecto
    root_row.cells[2].text = "Proyecto principal"
    root_row.cells[3].text = "—"
    for cell in root_row.cells:
        _set_cell_shading(cell, "1A3A5C")
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(10)
                run.font.name = "Calibri"
                run.font.color.rgb = _COLOR_BLANCO
        border_style = {"sz": "4", "val": "single", "color": "1A3A5C"}
        _set_cell_borders(
            cell, top=border_style, bottom=border_style,
            start=border_style, end=border_style,
        )

    # Agregar filas por cada fase y sus tareas
    for fase_idx, fase in enumerate(datos.fases):
        fase_codigo = str(fase_idx + 1)

        fase_row = table.add_row()
        fase_row.cells[0].text = fase_codigo
        fase_row.cells[1].text = fase.nombre
        fase_row.cells[2].text = "—"
        fase_row.cells[3].text = "—"

        color_hex = _COLORES_ETAPA[fase_idx % len(_COLORES_ETAPA)]
        for cell in fase_row.cells:
            _set_cell_shading(cell, color_hex)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(10)
                    run.font.name = "Calibri"
                    run.font.color.rgb = _COLOR_TEXTO
            border_style = {"sz": "4", "val": "single", "color": "BFBFBF"}
            _set_cell_borders(
                cell, top=border_style, bottom=border_style,
                start=border_style, end=border_style,
            )

        for tarea_idx, tarea in enumerate(fase.tareas, start=1):
            _agregar_filas_tabla_recursivo(
                table, tarea, fase_codigo, tarea_idx,
                nivel=2, color_etapa_idx=fase_idx,
            )

    doc.add_paragraph()
    _construir_tabla_firmas(doc)
    doc.add_page_break()


# ==============================================================================
# SECCIÓN 4 — MINUTA DE VALIDACIÓN (ANEXO INDEPENDIENTE)
# ==============================================================================

def _construir_seccion_minuta(doc, datos: DocumentoEDTInput, fecha: str) -> None:
    """Construye la Minuta de Validación como anexo independiente."""
    doc.add_page_break()

    # Encabezado corporativo de Minuta
    _construir_encabezado_plantilla(doc, datos, fecha, tipo_encabezado="Minuta")

    # Subtítulo "22.1 MINUTA EDT" con fondo azul claro
    tabla_subtitulo = doc.add_table(rows=1, cols=1)
    tabla_subtitulo.alignment = WD_TABLE_ALIGNMENT.CENTER
    tabla_subtitulo.autofit = True
    cell_sub = tabla_subtitulo.cell(0, 0)
    _set_cell_text(
        cell_sub, "22.1 MINUTA EDT",
        alignment=WD_ALIGN_PARAGRAPH.CENTER, font_size=14, bold=True,
        color=_COLOR_PRIMARIO,
    )
    _set_cell_shading(cell_sub, "EBF5FB")
    aplicar_bordes_uniformes(cell_sub, size="4", color="BFBFBF")
    doc.add_paragraph()

    # Objetivo de la reunión
    _add_styled_paragraph(
        doc, "Objetivo de la reunión",
        font_size=12, bold=True, color=_COLOR_PRIMARIO,
        alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=6,
    )

    _add_styled_paragraph(
        doc,
        "Validar la Estructura de Desglose del Trabajo (EDT) del proyecto, asegurando "
        "que todas las actividades, tiempos, recursos y responsables estén correctamente "
        "definidos y alineados con los objetivos del proyecto.",
        font_size=10, bold=False, color=_COLOR_TEXTO,
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=16,
    )

    # Tabla de actividades
    headers_act = ["Actividad", "Tiempo", "Recursos", "Responsable"]
    tabla_actividades = crear_tabla_con_encabezados(
        doc, headers_act, "0D2B4E",
        widths=[Cm(5), Cm(3), Cm(5), Cm(4.5)],
    )

    actividades = datos.actividades_minuta if datos.actividades_minuta else [
        ("Crear EDT", "1 día", "Plantillas aprobadas", "PM"),
        ("Validar EDT", "1 día", "Acta de constitución", "Director del proyecto"),
    ]

    for i, act in enumerate(actividades):
        if hasattr(act, 'actividad'):
            valores = [act.actividad, act.tiempo, act.recursos, act.responsable]
        else:
            valores = list(act)
        row = tabla_actividades.add_row()
        for j, valor in enumerate(valores):
            cell = row.cells[j]
            cell.text = str(valor)
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    run.font.size = Pt(9)
                    run.font.name = "Calibri"
                    run.font.color.rgb = _COLOR_TEXTO
            border_style = {"sz": "4", "val": "single", "color": "A9C4D8"}
            _set_cell_borders(
                cell, top=border_style, bottom=border_style,
                start=border_style, end=border_style,
            )

    doc.add_paragraph()

    # Conclusión
    _add_styled_paragraph(
        doc, "Conclusión",
        font_size=12, bold=True, color=_COLOR_PRIMARIO,
        alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=6,
    )

    _add_styled_paragraph(
        doc, datos.conclusion_minuta,
        font_size=10, bold=False, color=_COLOR_TEXTO,
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=16,
    )

    # Tabla de firmas
    _construir_tabla_firmas(doc)


# ==============================================================================
# FUNCIÓN PÚBLICA PRINCIPAL — GENERAR DOCUMENTO WORD
# ==============================================================================

def generar_documento_word(datos: DocumentoEDTInput) -> str:
    """
    Genera un documento corporativo .docx con la EDT completa del proyecto.

    El documento incluye tres secciones:
    1. DIAGRAMA EDT — Imagen PNG del diagrama visual generada vía mermaid.ink.
    2. TABLA BASE DE LA EDT — Tabla jerárquica con Código EDT, Nombre, Descripción
       Operativa e Hito para cada tarea.
    3. MINUTA DE VALIDACIÓN — Datos del proyecto y campos de firma para validación.
    """
    try:
        fecha = datos.fecha_generacion or datetime.now().strftime("%d/%m/%Y %H:%M:%S")

        if not datos.fases:
            return "❌ Error: No se proporcionaron fases. Revisa el JSON de entrada."

        # 1. Generar imagen del diagrama EDT
        try:
            image_bytes = generar_imagen_mermaid_para_docx(datos)
        except Exception:
            image_bytes = None

        # 2. Crear documento Word
        doc = Document()

        style = doc.styles["Normal"]
        font = style.font
        font.name = "Calibri"
        font.size = Pt(10)
        font.color.rgb = _COLOR_TEXTO

        for section in doc.sections:
            section.top_margin = Cm(2.54)
            section.bottom_margin = Cm(2.54)
            section.left_margin = Cm(2.54)
            section.right_margin = Cm(2.54)

        # 3. Encabezado plantilla corporativa
        _construir_encabezado_plantilla(doc, datos, fecha)

        # 4. Título centrado de la plantilla
        _add_styled_paragraph(
            doc, "PLANTILLA ESTRUCTURA DE DESGLOSE DEL TRABAJO (EDT)",
            font_size=18, bold=True, color=_COLOR_PRIMARIO,
            alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=4,
        )

        p_sep = doc.add_paragraph()
        run_sep = p_sep.add_run("─" * 80)
        run_sep.font.color.rgb = _COLOR_GRIS_CLARO
        run_sep.font.size = Pt(8)

        doc.add_paragraph()

        # 5. Diagrama EDT
        if image_bytes:
            _construir_seccion_diagrama(doc, image_bytes)
            doc.add_page_break()
        else:
            _add_styled_paragraph(
                doc,
                "⚠️ No se pudo generar la imagen del diagrama EDT. "
                "Verifique la conexión a internet para usar la API de mermaid.ink.",
                font_size=10, bold=False, color=RGBColor(0xDC, 0x35, 0x45),
                alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=12,
            )
            doc.add_page_break()

        _construir_seccion_tabla(doc, datos)
        _construir_seccion_minuta(doc, datos, fecha)

        # 6. Configurar encabezado y pie de página
        for section in doc.sections:
            header = section.header
            header.is_linked_to_previous = False
            header_p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
            header_p.text = f"{datos.nombre_proyecto}  |  EDT  |  {fecha}"
            header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            for run in header_p.runs:
                run.font.size = Pt(8)
                run.font.color.rgb = _COLOR_SECUNDARIO
                run.font.name = "Calibri"

            footer = section.footer
            footer.is_linked_to_previous = False
            footer_p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            footer_p.text = "Documento generado por MCP Gestión de Proyectos  |  Confidencial"
            footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in footer_p.runs:
                run.font.size = Pt(7)
                run.font.color.rgb = _COLOR_SECUNDARIO
                run.font.name = "Calibri"

        # 7. Guardar archivo
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        nombre_seguro = datos.nombre_proyecto.replace(" ", "_").replace("/", "-")
        nombre_archivo = f"{_PREFIJO_ARCHIVO}_{nombre_seguro}_{timestamp}.docx"
        ruta = _CARPETA_DOWNLOADS / nombre_archivo
        doc.save(str(ruta))

        return (
            f"✅ Documento corporativo generado exitosamente.\n"
            f"📁 Ruta: {ruta}\n"
            f"📄 Archivo: {nombre_archivo}\n\n"
            f"El documento contiene:\n"
            f"  1. Diagrama EDT visual {'(imagen embebida)' if image_bytes else '(no disponible — sin conexión)'}\n"
            f"  2. Tabla Base de la EDT jerarquizada\n"
            f"  3. Minuta de Validación con campos de firma"
        )

    except Exception as e:
        return (
            f"❌ Error al generar el documento: {str(e)}.\n"
            "Revisa los parámetros y vuelve a ejecutar la herramienta."
        )
