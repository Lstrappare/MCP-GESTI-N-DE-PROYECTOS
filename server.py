from __future__ import annotations

import base64
import httpx
from datetime import datetime
from pathlib import Path
from io import BytesIO

from mcp.server.fastmcp import FastMCP, Image
from pydantic import BaseModel, Field
from typing import List, Optional

# python-docx imports para generación de documentos Word
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

mcp = FastMCP("MCP-GestionProyectos")

# ==========================================
# MODELOS PYDANTIC — ESTRUCTURA RECURSIVA
# Una Tarea puede contener subtareas del
# mismo tipo Tarea (profundidad ilimitada).
# ==========================================

class Tarea(BaseModel):
    nombre: str = Field(..., min_length=1)
    subtareas: Optional[List[Tarea]] = []

# Necesario para que Pydantic resuelva la
# auto-referencia después de la definición.
Tarea.model_rebuild()


class Fase(BaseModel):
    nombre: str = Field(..., min_length=1)
    tareas: List[Tarea] = []


class EDTInput(BaseModel):
    nombre_proyecto: str = Field(..., min_length=1, description="El nombre oficial del proyecto.")
    fases: List[Fase] = Field(
        ...,
        description="""
        REGLA ESTRICTA PARA LA IA:
        1. Debes generar exactamente las etapas principales indicadas en los documentos fuente (ej. Inicio, Planeación, Ejecución, Control, Cierre).
        2. El anidamiento dentro de cada etapa debe ser COMPLETAMENTE VERTICAL Y LINEAL.
        3. Cada tarea padre debe tener un máximo de UNA (1) subtarea.
        4. Crea una cadena en cascada perfecta hacia abajo. ¡PROHIBIDO agrupar elementos horizontalmente!
        """
    )


# ==========================================
# FUNCIÓN AUXILIAR RECURSIVA
# Genera las conexiones Mermaid para un nodo
# y todos sus descendientes sin límite de
# profundidad. Usa un prefijo de ID único
# heredado por cada nivel de la llamada.
# ==========================================

def _agregar_nodos(
    tarea: Tarea,
    parent_id: str,
    node_id: str,
    lines: list[str],
) -> None:
    """Añade la arista parent→nodo y desciende recursivamente."""
    lines.append(f'    {parent_id} --> {node_id}["{tarea.nombre}"]')

    if tarea.subtareas:
        for idx, subtarea in enumerate(tarea.subtareas):
            child_id = f"{node_id}_{idx}"
            _agregar_nodos(subtarea, node_id, child_id, lines)


# ==========================================
# HELPER COMPARTIDO — CONSTRUCCIÓN MERMAID
# ==========================================

def _construir_mermaid(datos: EDTInput) -> str:
    """Construye y retorna el código Mermaid completo del EDT."""
    if not datos.fases:
        raise ValueError(
            "El JSON no contiene fases definidas. "
            "Revisa la extracción de los datos."
        )

    lines: list[str] = [
        "graph TD",
        f'    Root["{datos.nombre_proyecto}"]',
    ]

    for i, fase in enumerate(datos.fases):
        fase_id = f"F{i}"
        lines.append(f'    Root --> {fase_id}["{fase.nombre}"]')

        if not fase.tareas:
            raise ValueError(
                f"La fase '{fase.nombre}' vino sin tareas asignadas."
            )

        for j, tarea in enumerate(fase.tareas):
            tarea_id = f"T{i}_{j}"
            _agregar_nodos(tarea, fase_id, tarea_id, lines)

    # Estilos de color automáticos de MINTRANET
    lines.append("\n    %% Estilos automáticos de MINTRANET")
    lines.append("    style Root fill:#f8f9fa,stroke:#343a40,stroke-width:3px")

    colores = ["#d4edda", "#cce5ff", "#fff3cd", "#f8d7da", "#e2e3e5"]
    bordes  = ["#28a745", "#007bff", "#ffc107", "#dc3545", "#6c757d"]

    for i in range(len(datos.fases)):
        color_fondo = colores[i % len(colores)]
        color_borde = bordes[i % len(bordes)]
        lines.append(
            f"    style F{i} fill:{color_fondo},stroke:{color_borde},stroke-width:2px"
        )

    return "\n".join(lines)


# ==========================================
# HERRAMIENTAS MCP
# ==========================================

@mcp.tool()
def generar_edt(datos: EDTInput) -> str:
    """
    Transforma la estructura de un proyecto en un diagrama de Mermaid.

    INSTRUCCIONES CRÍTICAS PARA EL LLM ANTES DE INVOCAR ESTA HERRAMIENTA:
    - Lee todos los archivos .txt proporcionados por el usuario.
    - Extrae la información y fuérzala a una estructura de árbol de
      profundidad máxima pero de anchura mínima.
    - Para que el diagrama no se expanda hacia los lados, encadena cada
      tarea como hija única de la tarea anterior.
    - Solo invoca esta herramienta cuando hayas estructurado mentalmente
      el JSON cumpliendo la regla de 1 solo hijo por nodo.
    - Soporta profundidad de tareas ilimitada mediante recursividad.
    """
    try:
        return _construir_mermaid(datos)
    except Exception as e:
        # Retorno vital para que el Agente IA lea el error y corrija su JSON.
        return (
            f"Error de validación: {str(e)}. "
            "Corrige los parámetros y vuelve a ejecutar la herramienta."
        )


@mcp.tool()
def exportar_imagen_edt(datos: EDTInput) -> Image:
    """
    Genera el EDT en Mermaid, lo renderiza como imagen PNG y lo muestra
    directamente en el chat. También guarda el archivo en ~/Downloads/.

    Usa la API pública kroki.io (POST) para renderizar sin límite de URL,
    con fallback a mermaid.ink (GET) si kroki falla.

    CUÁNDO USAR ESTA HERRAMIENTA:
    - Cuando el usuario pida la imagen, el PNG, ver el diagrama o descargarlo.
    - Aplica las mismas reglas de estructura que generar_edt.

    ESTRUCTURA METODOLÓGICA FIJA — 5 ETAPAS OBLIGATORIAS:
    1. Etapa 0 (Inicio): Planeación Estratégica, Análisis del Entorno y Mercado,
       Estudio de Factibilidad, Definición de la Solución Tecnológica, Gestión Inicial del Proyecto.
    2. Etapa 1 (Planeación): Integración del Proyecto, Gestión de Interesados,
       Gestión del Alcance, Gestión de Requisitos.
    3. Etapa 2 (Ejecución): Tecnológica, Operativa, Recursos Humanos, Finanzas, Comercial.
    4. Etapa 3 (Control): Checklist (con hija 'Checklist de los 21 programas'),
       Control y seguimiento, Control de cambios.
    5. Etapa 4 (Cierre): Plantilla de resultados, Acta de cierre.

    Las tareas del proyecto deben anidarse como subtareas de nivel 3 o inferior
    dentro de esta estructura fija.

    REQUISITO DE NODOS:
    - Para garantizar una imagen legible y sin errores, se recomienda que el
      diagrama tenga exactamente 90 nodos (5 fases + 85 tareas).
    - El JSON debe estructurarse en 5 fases, cada una con una cadena vertical
      de 17 tareas anidadas secuencialmente (una subtarea por nodo).
    - Ejemplo: 5 fases × 17 tareas = 85 tareas + 5 fases = 90 nodos.
    """
    # 1. Construir el código Mermaid
    mermaid_code = _construir_mermaid(datos)

    # 2. Intentar con kroki.io (POST) – soporta diagramas grandes sin límite
    url_kroki = "https://kroki.io/mermaid/png"
    payload = {"diagram": mermaid_code}
    try:
        response = httpx.post(url_kroki, json=payload, timeout=30, follow_redirects=True)
        response.raise_for_status()
        image_bytes = response.content
    except Exception as e:
        # Fallback a mermaid.ink (GET) si kroki falla (puede dar 414 si es muy grande)
        encoded = base64.urlsafe_b64encode(mermaid_code.encode("utf-8")).decode("utf-8")
        url_mermaid = f"https://mermaid.ink/img/{encoded}?bgColor=white"
        response = httpx.get(url_mermaid, timeout=30, follow_redirects=True)
        response.raise_for_status()
        image_bytes = response.content

    # 3. Guardar en ~/Downloads/
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    nombre_seguro = datos.nombre_proyecto.replace(" ", "_").replace("/", "-")
    nombre_archivo = f"EDT_{nombre_seguro}_{timestamp}.png"
    ruta = Path.home() / "Downloads" / nombre_archivo
    ruta.write_bytes(image_bytes)

    # 4. Devolver objeto Image
    return Image(data=image_bytes, format="png")


# ==========================================
# MODELOS PYDANTIC — DOCUMENTO WORD EDT
# Modelos extendidos con campos adicionales
# para la Tabla Base (descripcion_operativa,
# hito). Independientes de los modelos
# originales para mantener retrocompatibilidad.
# ==========================================

class TareaEDT(BaseModel):
    nombre: str = Field(..., min_length=1)
    descripcion_operativa: Optional[str] = "N/A"
    hito: Optional[str] = "N/A"
    subtareas: Optional[List["TareaEDT"]] = []

TareaEDT.model_rebuild()


class FaseEDT(BaseModel):
    nombre: str = Field(..., min_length=1)
    tareas: List[TareaEDT] = []


class DocumentoEDTInput(BaseModel):
    nombre_proyecto: str = Field(..., min_length=1, description="Nombre oficial del proyecto.")
    id_proyecto: Optional[str] = Field("N/A", description="Identificador único del proyecto.")
    presupuesto_total: Optional[str] = Field("N/A", description="Presupuesto total del proyecto.")
    fecha_generacion: Optional[str] = Field(
        None,
        description="Fecha de generación del documento. Se auto-genera si no se proporciona.",
    )
    fases: List[FaseEDT] = Field(
        ...,
        description="""
        ESTRUCTURA METODOLÓGICA FIJA — 5 ETAPAS OBLIGATORIAS:
        1. Etapa 0 (Inicio): Planeación Estratégica, Análisis del Entorno y Mercado,
           Estudio de Factibilidad, Definición de la Solución Tecnológica, Gestión Inicial del Proyecto.
        2. Etapa 1 (Planeación): Integración del Proyecto, Gestión de Interesados,
           Gestión del Alcance, Gestión de Requisitos.
        3. Etapa 2 (Ejecución): Tecnológica, Operativa, Recursos Humanos, Finanzas, Comercial.
        4. Etapa 3 (Control): Checklist (con hija 'Checklist de los 21 programas'),
           Control y seguimiento, Control de cambios.
        5. Etapa 4 (Cierre): Plantilla de resultados, Acta de cierre.

        Las tareas del proyecto deben anidarse como subtareas de nivel 3 o inferior
        dentro de esta estructura fija.
        """,
    )


# ==========================================
# HELPERS — GENERACIÓN DE DOCUMENTO WORD
# ==========================================

# Paleta de colores corporativos
_COLOR_PRIMARIO = RGBColor(0x1B, 0x3A, 0x5C)       # Azul oscuro corporativo
_COLOR_SECUNDARIO = RGBColor(0x2E, 0x86, 0xC1)      # Azul medio
_COLOR_ACENTO = RGBColor(0x17, 0xA5, 0x89)           # Verde-azul
_COLOR_TEXTO = RGBColor(0x2C, 0x3E, 0x50)            # Gris oscuro
_COLOR_GRIS_CLARO = RGBColor(0xEC, 0xF0, 0xF1)       # Fondo gris claro
_COLOR_BLANCO = RGBColor(0xFF, 0xFF, 0xFF)

# Colores de fondo para cada etapa (hex sin #)
_COLORES_ETAPA = [
    "D4EDDA",  # Verde claro — Inicio
    "CCE5FF",  # Azul claro — Planeación
    "FFF3CD",  # Amarillo claro — Ejecución
    "F8D7DA",  # Rojo claro — Control
    "E2E3E5",  # Gris claro — Cierre
]

# Colores de fondo para subtareas base (nivel 2)
_COLOR_SUBTAREA_BASE = "F0F4F8"


def _set_cell_shading(cell, color_hex: str) -> None:
    """Aplica color de fondo a una celda de tabla Word."""
    shading_elm = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>'
    )
    cell._tc.get_or_add_tcPr().append(shading_elm)


def _set_cell_borders(cell, **kwargs) -> None:
    """
    Aplica bordes a una celda. Parámetros: top, bottom, start, end.
    Cada uno recibe un dict con: sz (tamaño), val (tipo), color.
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')

    for edge, attrs in kwargs.items():
        element = parse_xml(
            f'<w:{edge} {nsdecls("w")} '
            f'w:sz="{attrs.get("sz", 4)}" '
            f'w:val="{attrs.get("val", "single")}" '
            f'w:color="{attrs.get("color", "BFBFBF")}"/>'
        )
        tcBorders.append(element)

    tcPr.append(tcBorders)


def _format_paragraph(paragraph, font_size=10, bold=False, color=None, alignment=None, font_name="Calibri"):
    """Formatea un párrafo con estilos corporativos."""
    if alignment is not None:
        paragraph.alignment = alignment
    for run in paragraph.runs:
        run.font.size = Pt(font_size)
        run.font.bold = bold
        run.font.name = font_name
        if color:
            run.font.color.rgb = color


def _add_styled_paragraph(doc, text, font_size=10, bold=False, color=None, alignment=None, space_after=6):
    """Agrega un párrafo con estilo al documento."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.name = "Calibri"
    if color:
        run.font.color.rgb = color
    if alignment is not None:
        p.alignment = alignment
    p.paragraph_format.space_after = Pt(space_after)
    return p


def _convertir_a_edt_input(datos: DocumentoEDTInput) -> EDTInput:
    """
    Convierte DocumentoEDTInput a EDTInput para reutilizar
    la lógica de construcción Mermaid existente.
    """
    def _convertir_tarea(t: TareaEDT) -> Tarea:
        return Tarea(
            nombre=t.nombre,
            subtareas=[_convertir_tarea(st) for st in (t.subtareas or [])],
        )

    fases_convertidas = []
    for fase in datos.fases:
        fases_convertidas.append(
            Fase(
                nombre=fase.nombre,
                tareas=[_convertir_tarea(t) for t in fase.tareas],
            )
        )

    return EDTInput(
        nombre_proyecto=datos.nombre_proyecto,
        fases=fases_convertidas,
    )


def _generar_imagen_mermaid(datos: DocumentoEDTInput) -> bytes:
    """Genera la imagen PNG del diagrama EDT usando mermaid.ink."""
    edt_input = _convertir_a_edt_input(datos)
    mermaid_code = _construir_mermaid(edt_input)
    encoded = base64.urlsafe_b64encode(mermaid_code.encode("utf-8")).decode("utf-8")
    url_imagen = f"https://mermaid.ink/img/{encoded}?bgColor=white"
    response = httpx.get(url_imagen, timeout=30, follow_redirects=True)
    response.raise_for_status()
    return response.content


def _agregar_filas_tabla_recursivo(
    table, tarea: TareaEDT, codigo_padre: str, indice: int,
    nivel: int, color_etapa_idx: int
) -> int:
    """
    Agrega filas recursivamente a la tabla de la EDT.
    Retorna el número total de tareas procesadas (para mantener conteo).
    """
    codigo = f"{codigo_padre}.{indice}"
    row = table.add_row()
    cells = row.cells

    # Determinar indentación visual según nivel
    indent = "    " * (nivel - 2) if nivel >= 2 else ""
    nombre_display = f"{indent}{tarea.nombre}" if nivel >= 3 else tarea.nombre

    cells[0].text = codigo
    cells[1].text = nombre_display
    cells[2].text = tarea.descripcion_operativa or "N/A"
    cells[3].text = tarea.hito or "N/A"

    # Estilizar celdas
    for cell in cells:
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in paragraph.runs:
                run.font.size = Pt(9)
                run.font.name = "Calibri"
                run.font.color.rgb = _COLOR_TEXTO

    # Aplicar colores según nivel
    if nivel == 1:
        # Nivel de etapa — color de etapa
        color_hex = _COLORES_ETAPA[color_etapa_idx % len(_COLORES_ETAPA)]
        for cell in cells:
            _set_cell_shading(cell, color_hex)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(10)
    elif nivel == 2:
        # Subtarea base — fondo sutil
        for cell in cells:
            _set_cell_shading(cell, _COLOR_SUBTAREA_BASE)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(9)

    # Aplicar bordes sutiles
    border_style = {"sz": "4", "val": "single", "color": "BFBFBF"}
    for cell in cells:
        _set_cell_borders(
            cell, top=border_style, bottom=border_style,
            start=border_style, end=border_style,
        )

    # Procesar subtareas recursivamente
    count = 1
    if tarea.subtareas:
        for sub_idx, subtarea in enumerate(tarea.subtareas, start=1):
            count += _agregar_filas_tabla_recursivo(
                table, subtarea, codigo, sub_idx,
                nivel + 1, color_etapa_idx,
            )

    return count


def _construir_portada(doc, datos: DocumentoEDTInput, fecha: str) -> None:
    """Construye la portada corporativa del documento."""
    # Espaciado superior
    for _ in range(4):
        doc.add_paragraph()

    # Línea decorativa superior
    p_line = doc.add_paragraph()
    p_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_line.add_run("━" * 50)
    run.font.color.rgb = _COLOR_SECUNDARIO
    run.font.size = Pt(14)

    # Título del documento
    _add_styled_paragraph(
        doc, "ESTRUCTURA DE DESGLOSE DEL TRABAJO",
        font_size=26, bold=True, color=_COLOR_PRIMARIO,
        alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=4,
    )
    _add_styled_paragraph(
        doc, "(EDT / WBS)",
        font_size=16, bold=False, color=_COLOR_SECUNDARIO,
        alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=20,
    )

    # Línea decorativa
    p_line2 = doc.add_paragraph()
    p_line2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p_line2.add_run("━" * 50)
    run2.font.color.rgb = _COLOR_SECUNDARIO
    run2.font.size = Pt(14)

    doc.add_paragraph()

    # Datos del proyecto en tabla de portada
    tabla_info = doc.add_table(rows=3, cols=2)
    tabla_info.alignment = WD_TABLE_ALIGNMENT.CENTER

    datos_portada = [
        ("Proyecto:", datos.nombre_proyecto),
        ("ID del Proyecto:", datos.id_proyecto or "N/A"),
        ("Presupuesto Total:", datos.presupuesto_total or "N/A"),
    ]

    for i, (label, valor) in enumerate(datos_portada):
        cell_label = tabla_info.cell(i, 0)
        cell_valor = tabla_info.cell(i, 1)
        cell_label.text = label
        cell_valor.text = valor

        for paragraph in cell_label.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(12)
                run.font.name = "Calibri"
                run.font.color.rgb = _COLOR_PRIMARIO

        for paragraph in cell_valor.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in paragraph.runs:
                run.font.size = Pt(12)
                run.font.name = "Calibri"
                run.font.color.rgb = _COLOR_TEXTO

        # Remover bordes de la tabla de portada
        for cell in [cell_label, cell_valor]:
            _set_cell_borders(
                cell,
                top={"sz": "0", "val": "none", "color": "FFFFFF"},
                bottom={"sz": "0", "val": "none", "color": "FFFFFF"},
                start={"sz": "0", "val": "none", "color": "FFFFFF"},
                end={"sz": "0", "val": "none", "color": "FFFFFF"},
            )

    doc.add_paragraph()

    # Fecha
    _add_styled_paragraph(
        doc, f"Fecha de generación: {fecha}",
        font_size=11, bold=False, color=_COLOR_SECUNDARIO,
        alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=10,
    )

    # Línea decorativa inferior
    p_line3 = doc.add_paragraph()
    p_line3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p_line3.add_run("━" * 50)
    run3.font.color.rgb = _COLOR_SECUNDARIO
    run3.font.size = Pt(14)

    # Salto de página
    doc.add_page_break()


def _construir_seccion_diagrama(doc, image_bytes: bytes) -> None:
    """Inserta la sección del diagrama EDT como imagen."""
    _add_styled_paragraph(
        doc, "1. DIAGRAMA EDT",
        font_size=16, bold=True, color=_COLOR_PRIMARIO,
        alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=6,
    )

    # Línea separadora
    p_sep = doc.add_paragraph()
    run_sep = p_sep.add_run("─" * 80)
    run_sep.font.color.rgb = _COLOR_GRIS_CLARO
    run_sep.font.size = Pt(8)

    _add_styled_paragraph(
        doc,
        "El siguiente diagrama presenta la Estructura de Desglose del Trabajo "
        "del proyecto de forma visual y jerárquica.",
        font_size=10, bold=False, color=_COLOR_TEXTO,
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=12,
    )

    # Insertar imagen
    image_stream = BytesIO(image_bytes)
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_img = p_img.add_run()
    run_img.add_picture(image_stream, width=Inches(6.5))

    doc.add_paragraph()
    doc.add_page_break()


def _construir_seccion_tabla(doc, datos: DocumentoEDTInput) -> None:
    """Construye la Tabla Base de la EDT jerarquizada."""
    _add_styled_paragraph(
        doc, "2. TABLA BASE DE LA EDT",
        font_size=16, bold=True, color=_COLOR_PRIMARIO,
        alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=6,
    )

    p_sep = doc.add_paragraph()
    run_sep = p_sep.add_run("─" * 80)
    run_sep.font.color.rgb = _COLOR_GRIS_CLARO
    run_sep.font.size = Pt(8)

    _add_styled_paragraph(
        doc,
        "La siguiente tabla presenta la codificación EDT jerárquica, la descripción "
        "operativa de cada tarea y sus hitos asociados.",
        font_size=10, bold=False, color=_COLOR_TEXTO,
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=12,
    )

    # Crear tabla con encabezados
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    # Encabezados
    headers = ["Código EDT", "Nombre de la Tarea", "Descripción Operativa", "Hito"]
    header_row = table.rows[0]
    for i, header in enumerate(headers):
        cell = header_row.cells[i]
        cell.text = header
        _set_cell_shading(cell, "1B3A5C")  # Azul oscuro
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(10)
                run.font.name = "Calibri"
                run.font.color.rgb = _COLOR_BLANCO
        border_style = {"sz": "6", "val": "single", "color": "1B3A5C"}
        _set_cell_borders(
            cell, top=border_style, bottom=border_style,
            start=border_style, end=border_style,
        )

    # Ajustar anchos de columnas
    widths = [Cm(2.5), Cm(6), Cm(6.5), Cm(2.5)]
    for i, width in enumerate(widths):
        header_row.cells[i].width = width

    # Agregar fila del proyecto raíz
    root_row = table.add_row()
    root_row.cells[0].text = "0"
    root_row.cells[1].text = datos.nombre_proyecto
    root_row.cells[2].text = "Proyecto principal"
    root_row.cells[3].text = "—"
    for cell in root_row.cells:
        _set_cell_shading(cell, "34495E")
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(10)
                run.font.name = "Calibri"
                run.font.color.rgb = _COLOR_BLANCO
        border_style = {"sz": "4", "val": "single", "color": "34495E"}
        _set_cell_borders(
            cell, top=border_style, bottom=border_style,
            start=border_style, end=border_style,
        )

    # Agregar filas por cada fase y sus tareas
    for fase_idx, fase in enumerate(datos.fases):
        fase_codigo = str(fase_idx + 1)

        # Fila de la fase (nivel 0 de la tabla = nivel 1 del árbol)
        fase_row = table.add_row()
        fase_row.cells[0].text = fase_codigo
        fase_row.cells[1].text = fase.nombre
        fase_row.cells[2].text = "—"
        fase_row.cells[3].text = "—"

        color_hex = _COLORES_ETAPA[fase_idx % len(_COLORES_ETAPA)]
        for cell in fase_row.cells:
            _set_cell_shading(cell, color_hex)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(10)
                    run.font.name = "Calibri"
                    run.font.color.rgb = _COLOR_TEXTO
            border_style = {"sz": "4", "val": "single", "color": "BFBFBF"}
            _set_cell_borders(
                cell, top=border_style, bottom=border_style,
                start=border_style, end=border_style,
            )

        # Tareas de la fase
        for tarea_idx, tarea in enumerate(fase.tareas, start=1):
            _agregar_filas_tabla_recursivo(
                table, tarea, fase_codigo, tarea_idx,
                nivel=2, color_etapa_idx=fase_idx,
            )

    doc.add_paragraph()
    doc.add_page_break()


def _construir_seccion_minuta(doc, datos: DocumentoEDTInput, fecha: str) -> None:
    """Construye la Minuta de Validación."""
    _add_styled_paragraph(
        doc, "3. MINUTA DE VALIDACIÓN",
        font_size=16, bold=True, color=_COLOR_PRIMARIO,
        alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=6,
    )

    p_sep = doc.add_paragraph()
    run_sep = p_sep.add_run("─" * 80)
    run_sep.font.color.rgb = _COLOR_GRIS_CLARO
    run_sep.font.size = Pt(8)

    _add_styled_paragraph(
        doc,
        "La presente minuta certifica la validación y aprobación de la Estructura "
        "de Desglose del Trabajo (EDT) del proyecto.",
        font_size=10, bold=False, color=_COLOR_TEXTO,
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=16,
    )

    # Tabla de datos del proyecto
    _add_styled_paragraph(
        doc, "Datos del Proyecto",
        font_size=13, bold=True, color=_COLOR_SECUNDARIO,
        alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=8,
    )

    tabla_datos = doc.add_table(rows=4, cols=2)
    tabla_datos.alignment = WD_TABLE_ALIGNMENT.LEFT

    datos_proyecto = [
        ("Nombre del Proyecto", datos.nombre_proyecto),
        ("ID del Proyecto", datos.id_proyecto or "N/A"),
        ("Presupuesto Total", datos.presupuesto_total or "N/A"),
        ("Fecha de Generación", fecha),
    ]

    for i, (label, valor) in enumerate(datos_proyecto):
        cell_label = tabla_datos.cell(i, 0)
        cell_valor = tabla_datos.cell(i, 1)
        cell_label.text = label
        cell_valor.text = valor

        _set_cell_shading(cell_label, "EBF5FB")
        border_style = {"sz": "4", "val": "single", "color": "BDC3C7"}

        for cell in [cell_label, cell_valor]:
            _set_cell_borders(
                cell, top=border_style, bottom=border_style,
                start=border_style, end=border_style,
            )
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)
                    run.font.name = "Calibri"
                    run.font.color.rgb = _COLOR_TEXTO

        for paragraph in cell_label.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True

    # Ajustar anchos
    for row in tabla_datos.rows:
        row.cells[0].width = Cm(6)
        row.cells[1].width = Cm(11)

    doc.add_paragraph()

    # Total de tareas
    total_tareas = 0
    for fase in datos.fases:
        def _contar(tareas):
            count = 0
            for t in tareas:
                count += 1
                if t.subtareas:
                    count += _contar(t.subtareas)
            return count
        total_tareas += _contar(fase.tareas)

    _add_styled_paragraph(
        doc, f"Total de tareas/actividades identificadas: {total_tareas}",
        font_size=10, bold=True, color=_COLOR_TEXTO,
        alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=6,
    )

    _add_styled_paragraph(
        doc, f"Total de etapas: {len(datos.fases)}",
        font_size=10, bold=True, color=_COLOR_TEXTO,
        alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=16,
    )

    # Tabla de firmas
    _add_styled_paragraph(
        doc, "Firmas de Validación",
        font_size=13, bold=True, color=_COLOR_SECUNDARIO,
        alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=12,
    )

    tabla_firmas = doc.add_table(rows=4, cols=3)
    tabla_firmas.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Encabezados de firma
    headers_firma = ["Rol", "Nombre Completo", "Firma"]
    for i, header in enumerate(headers_firma):
        cell = tabla_firmas.cell(0, i)
        cell.text = header
        _set_cell_shading(cell, "1B3A5C")
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(10)
                run.font.name = "Calibri"
                run.font.color.rgb = _COLOR_BLANCO
        border_style = {"sz": "6", "val": "single", "color": "1B3A5C"}
        _set_cell_borders(
            cell, top=border_style, bottom=border_style,
            start=border_style, end=border_style,
        )

    roles = ["Project Manager", "Sponsor / Patrocinador", "Líder Técnico"]
    for i, rol in enumerate(roles, start=1):
        tabla_firmas.cell(i, 0).text = rol
        tabla_firmas.cell(i, 1).text = ""
        tabla_firmas.cell(i, 2).text = ""

        for j in range(3):
            cell = tabla_firmas.cell(i, j)
            border_style = {"sz": "4", "val": "single", "color": "BDC3C7"}
            _set_cell_borders(
                cell, top=border_style, bottom=border_style,
                start=border_style, end=border_style,
            )
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)
                    run.font.name = "Calibri"

            # Altura mínima para espacio de firma
            if j == 2:
                cell.width = Cm(5)

        # Estilizar la celda del rol
        for paragraph in tabla_firmas.cell(i, 0).paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = _COLOR_PRIMARIO

    # Ajustar anchos de columnas de firma
    for row in tabla_firmas.rows:
        row.cells[0].width = Cm(5)
        row.cells[1].width = Cm(7)
        row.cells[2].width = Cm(5)

    doc.add_paragraph()
    doc.add_paragraph()

    # Nota al pie
    _add_styled_paragraph(
        doc,
        "Este documento fue generado automáticamente por el sistema MCP de Gestión "
        "de Proyectos. La validación de la EDT requiere la firma de los responsables "
        "indicados anteriormente.",
        font_size=8, bold=False, color=_COLOR_SECUNDARIO,
        alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=4,
    )


# ==========================================
# HERRAMIENTA MCP — DOCUMENTO WORD
# ==========================================

@mcp.tool()
def generar_documento_proyecto_word(datos: DocumentoEDTInput) -> str:
    """
    Genera un documento corporativo .docx con la EDT completa del proyecto.

    El documento incluye tres secciones:
    1. DIAGRAMA EDT — Imagen PNG del diagrama visual generada vía mermaid.ink.
    2. TABLA BASE DE LA EDT — Tabla jerárquica con Código EDT, Nombre, Descripción
       Operativa e Hito para cada tarea.
    3. MINUTA DE VALIDACIÓN — Datos del proyecto y campos de firma para validación.

    INSTRUCCIONES PARA EL LLM:
    - Lee todos los archivos .txt del usuario y extrae: nombre del proyecto,
      ID del proyecto, presupuesto total, y todas las tareas/actividades.
    - Clasifica TODAS las tareas dentro de la Estructura Metodológica Fija de 5 etapas.
    - Para cada tarea, proporciona: descripcion_operativa (breve descripción) y
      hito (ej. "Hito 1", "N/A").
    - Las tareas del proyecto deben anidarse en nivel 3 o inferior.
    - El archivo .docx se guarda automáticamente en ~/Downloads/.
    """
    try:
        # Fecha de generación
        fecha = datos.fecha_generacion or datetime.now().strftime("%d/%m/%Y %H:%M:%S")

        # Validar que hay fases
        if not datos.fases:
            return "❌ Error: No se proporcionaron fases. Revisa el JSON de entrada."

        # 1. Generar imagen del diagrama EDT
        try:
            image_bytes = _generar_imagen_mermaid(datos)
        except Exception as img_err:
            image_bytes = None

        # 2. Crear documento Word
        doc = Document()

        # Configurar estilos globales del documento
        style = doc.styles["Normal"]
        font = style.font
        font.name = "Calibri"
        font.size = Pt(10)
        font.color.rgb = _COLOR_TEXTO

        # Configurar márgenes
        for section in doc.sections:
            section.top_margin = Cm(2.54)
            section.bottom_margin = Cm(2.54)
            section.left_margin = Cm(2.54)
            section.right_margin = Cm(2.54)

        # 3. Construir secciones
        _construir_portada(doc, datos, fecha)

        if image_bytes:
            _construir_seccion_diagrama(doc, image_bytes)
        else:
            _add_styled_paragraph(
                doc, "1. DIAGRAMA EDT",
                font_size=16, bold=True, color=_COLOR_PRIMARIO,
                alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=6,
            )
            _add_styled_paragraph(
                doc,
                "⚠️ No se pudo generar la imagen del diagrama EDT. "
                "Verifique la conexión a internet para usar la API de mermaid.ink.",
                font_size=10, bold=False, color=RGBColor(0xDC, 0x35, 0x45),
                alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=12,
            )
            doc.add_page_break()

        _construir_seccion_tabla(doc, datos)
        _construir_seccion_minuta(doc, datos, fecha)

        # 4. Configurar encabezado y pie de página
        for section in doc.sections:
            # Encabezado
            header = section.header
            header.is_linked_to_previous = False
            header_p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
            header_p.text = f"{datos.nombre_proyecto}  |  EDT  |  {fecha}"
            header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            for run in header_p.runs:
                run.font.size = Pt(8)
                run.font.color.rgb = _COLOR_SECUNDARIO
                run.font.name = "Calibri"

            # Pie de página
            footer = section.footer
            footer.is_linked_to_previous = False
            footer_p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            footer_p.text = "Documento generado por MCP Gestión de Proyectos  |  Confidencial"
            footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in footer_p.runs:
                run.font.size = Pt(7)
                run.font.color.rgb = _COLOR_GRIS_CLARO
                run.font.name = "Calibri"

        # 5. Guardar archivo
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        nombre_seguro = datos.nombre_proyecto.replace(" ", "_").replace("/", "-")
        nombre_archivo = f"EDT_{nombre_seguro}_{timestamp}.docx"
        ruta = Path.home() / "Downloads" / nombre_archivo
        doc.save(str(ruta))

        return (
            f"✅ Documento corporativo generado exitosamente.\n"
            f"📁 Ruta: {ruta}\n"
            f"📄 Archivo: {nombre_archivo}\n\n"
            f"El documento contiene:\n"
            f"  1. Diagrama EDT visual {'(imagen embebida)' if image_bytes else '(no disponible — sin conexión)'}\n"
            f"  2. Tabla Base de la EDT jerarquizada\n"
            f"  3. Minuta de Validación con campos de firma"
        )

    except Exception as e:
        return (
            f"❌ Error al generar el documento: {str(e)}.\n"
            "Revisa los parámetros y vuelve a ejecutar la herramienta."
        )


if __name__ == "__main__":
    print("Iniciando Servidor MCP de Gestión de Proyectos...")
    mcp.run(transport='stdio')